import io
from pathlib import Path

import pdfplumber
import pandas as pd
from docx import Document


def read_pdf_meta(file_path: str) -> dict:
    """Trả về metadata + số trang + page size + preview nội dung đầu (5000 ký tự)."""
    with pdfplumber.open(file_path) as pdf:
        raw_meta = pdf.metadata or {}
        total_pages = len(pdf.pages)
        first_page = pdf.pages[0] if total_pages > 0 else None
        page_size = {
            "width": first_page.width if first_page else None,
            "height": first_page.height if first_page else None,
        }

        preview_parts = []
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text() or ""
            if text:
                preview_parts.append(f"--- Trang {i} ---\n{text}")
            if len("\n\n".join(preview_parts)) >= 5000:
                break

        preview = "\n\n".join(preview_parts)[:5000]

    def _clean(v):
        return v.strip() if isinstance(v, str) else v

    return {
        "total_pages": total_pages,
        "page_size": page_size,
        "metadata": {
            "title":        _clean(raw_meta.get("Title", "")),
            "author":       _clean(raw_meta.get("Author", "")),
            "subject":      _clean(raw_meta.get("Subject", "")),
            "keywords":     _clean(raw_meta.get("Keywords", "")),
            "creator":      _clean(raw_meta.get("Creator", "")),
            "producer":     _clean(raw_meta.get("Producer", "")),
            "created_at":   _clean(raw_meta.get("CreationDate", "")),
            "modified_at":  _clean(raw_meta.get("ModDate", "")),
        },
        "preview": preview,
    }


def read_pdf(file_path: str) -> str:
    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text() or ""
            tables = page.extract_tables()
            if text:
                text_parts.append(f"--- Trang {i} ---\n{text}")
            for table in tables:
                if table:
                    df = pd.DataFrame(table[1:], columns=table[0])
                    text_parts.append(f"[Bảng trang {i}]\n{df.to_markdown(index=False)}")
    return "\n\n".join(text_parts)


def _page_range(total: int, page_start: int, page_end: int | None) -> range:
    start = max(1, page_start) - 1
    end = min(total, page_end or total) - 1
    return range(start, end + 1)


def _extract_images(page) -> list[dict]:
    return [
        {
            "x0": img["x0"], "y0": img["y0"],
            "x1": img["x1"], "y1": img["y1"],
            "width": img["width"], "height": img["height"],
        }
        for img in page.images
    ]


def _extract_annots(page) -> list[dict]:
    result = []
    for a in page.annots:
        result.append({
            "uri":  a.get("uri"),
            "title": a.get("title"),
            "x0": a.get("x0"), "y0": a.get("y0"),
            "x1": a.get("x1"), "y1": a.get("y1"),
        })
    return result


def read_pdf_pages(file_path: str, page_start: int = 1, page_end: int | None = None) -> list[dict]:
    """
    Trả về danh sách các trang theo khoảng [page_start, page_end] (1-indexed, inclusive).
    Mỗi phần tử: {page_number, text, tables, images, annots}
    """
    pages = []
    with pdfplumber.open(file_path) as pdf:
        total = len(pdf.pages)
        for i in _page_range(total, page_start, page_end):
            page = pdf.pages[i]
            tables = []
            for tbl in page.extract_tables():
                if tbl:
                    df = pd.DataFrame(tbl[1:], columns=tbl[0])
                    tables.append(df.to_markdown(index=False))
            pages.append({
                "page_number": i + 1,
                "text": page.extract_text() or "",
                "tables": tables,
                "images": _extract_images(page),
                "annots": _extract_annots(page),
            })
    return pages


def read_pdf_pages_detailed(file_path: str, page_start: int = 1, page_end: int | None = None) -> list[dict]:
    """
    Giống read_pdf_pages nhưng thêm thông tin vector chi tiết:
    chars (font, size, màu, tọa độ), lines, rects.
    """
    pages = []
    with pdfplumber.open(file_path) as pdf:
        total = len(pdf.pages)
        for i in _page_range(total, page_start, page_end):
            page = pdf.pages[i]
            tables = []
            for tbl in page.extract_tables():
                if tbl:
                    df = pd.DataFrame(tbl[1:], columns=tbl[0])
                    tables.append(df.to_markdown(index=False))

            chars = [
                {
                    "text": c["text"],
                    "fontname": c["fontname"],
                    "size": round(c["size"], 2),
                    "x0": round(c["x0"], 2), "y0": round(c["y0"], 2),
                    "x1": round(c["x1"], 2), "y1": round(c["y1"], 2),
                    "color": c.get("non_stroking_color"),
                    "upright": c["upright"],
                }
                for c in page.chars
            ]

            lines = [
                {
                    "x0": round(l["x0"], 2), "y0": round(l["y0"], 2),
                    "x1": round(l["x1"], 2), "y1": round(l["y1"], 2),
                    "width": round(l.get("width", 0), 2),
                }
                for l in page.lines
            ]

            rects = [
                {
                    "x0": round(r["x0"], 2), "y0": round(r["y0"], 2),
                    "x1": round(r["x1"], 2), "y1": round(r["y1"], 2),
                    "width": round(r.get("width", 0), 2),
                    "height": round(r.get("height", 0), 2),
                }
                for r in page.rects
            ]

            pages.append({
                "page_number": i + 1,
                "width": page.width,
                "height": page.height,
                "rotation": page.rotation,
                "text": page.extract_text() or "",
                "tables": tables,
                "images": _extract_images(page),
                "annots": _extract_annots(page),
                "chars": chars,
                "lines": lines,
                "rects": rects,
            })
    return pages


def get_pdf_page_count(file_path: str) -> int:
    with pdfplumber.open(file_path) as pdf:
        return len(pdf.pages)


def read_pdf_tables_custom(file_path: str, page_number: int, strategy: str = "text") -> dict:
    """Extract tables from one page using a custom detection strategy.
    strategy: 'lines' (needs visible borders) | 'text' (word-position based, no borders needed)
    """
    if strategy not in {"lines", "text"}:
        strategy = "text"
    settings = {
        "vertical_strategy": strategy,
        "horizontal_strategy": strategy,
        "snap_tolerance": 5,
        "join_tolerance": 5,
        "edge_min_length": 3,
        "min_words_vertical": 3,
        "min_words_horizontal": 1,
    }
    with pdfplumber.open(file_path) as pdf:
        total = len(pdf.pages)
        if page_number < 1 or page_number > total:
            return {"page_number": page_number, "strategy": strategy, "tables": [], "table_count": 0}
        page = pdf.pages[page_number - 1]
        tables = []
        for tbl in page.extract_tables(settings):
            if tbl:
                try:
                    df = pd.DataFrame(tbl[1:], columns=tbl[0])
                    tables.append(df.to_markdown(index=False))
                except Exception:
                    tables.append(str(tbl))
    return {"page_number": page_number, "strategy": strategy, "tables": tables, "table_count": len(tables)}


def extract_pdf_structure(file_path: str) -> dict:
    """Scan all pages and return a heading hierarchy inferred from font sizes.
    Returns: {body_size, heading_sizes, total_headings, headings: [{level, text, page, size, bold}]}
    """
    from collections import Counter

    all_lines: list[dict] = []
    size_counter: Counter = Counter()

    with pdfplumber.open(file_path) as pdf:
        for page_i, page in enumerate(pdf.pages, 1):
            if not page.chars:
                continue
            lines_dict: dict[int, list] = {}
            for c in page.chars:
                if not c.get("text", "").strip():
                    continue
                y_key = round(c["y0"] / 2) * 2
                lines_dict.setdefault(y_key, []).append(c)

            for y_key in sorted(lines_dict, reverse=True):
                chars = sorted(lines_dict[y_key], key=lambda c: c["x0"])
                text = "".join(c["text"] for c in chars).strip()
                if not text:
                    continue
                sizes = [c["size"] for c in chars if c.get("size")]
                if not sizes:
                    continue
                max_size = max(sizes)
                is_bold = any(
                    "bold" in c.get("fontname", "").lower() for c in chars
                )
                size_counter[round(max_size)] += 1
                all_lines.append({"page": page_i, "size": round(max_size, 1), "bold": is_bold, "text": text})

    if not size_counter:
        return {"error": "No text found"}

    body_size = size_counter.most_common(1)[0][0]
    heading_threshold = body_size * 1.15
    heading_sizes = sorted(
        {round(l["size"]) for l in all_lines if l["size"] >= heading_threshold},
        reverse=True,
    )
    size_to_level = {s: i + 1 for i, s in enumerate(heading_sizes[:4])}

    headings = []
    for line in all_lines:
        rounded = round(line["size"])
        if rounded in size_to_level:
            headings.append({
                "level": size_to_level[rounded],
                "text": line["text"][:150],
                "page": line["page"],
                "size": line["size"],
                "bold": line["bold"],
            })

    return {
        "body_size": body_size,
        "heading_sizes": heading_sizes,
        "total_headings": len(headings),
        "headings": headings,
    }


def read_csv(file_path: str) -> str:
    df = pd.read_csv(file_path)
    return df.to_markdown(index=False)


def read_tabular_meta(file_path: str) -> dict:
    """Trả về metadata + raw preview 6 dòng đầu (header=None) cho CSV / XLSX / XLS.
    LLM sẽ tự xác định header row từ raw_preview."""
    ext = Path(file_path).suffix.lower()
    result = {"file_type": "xlsx", "sheets": []}

    if ext == ".csv":
        df = pd.read_csv(file_path)
        result["sheets"] = [_sheet_meta("(csv)", df)]
    else:
        xl = pd.ExcelFile(file_path)
        result["sheet_names"] = xl.sheet_names
        for name in xl.sheet_names:
            raw = xl.parse(name, header=None, nrows=6)
            raw_preview = raw.fillna("").astype(str).values.tolist()
            total_rows = xl.parse(name, header=None).shape[0]
            result["sheets"].append({
                "sheet": name,
                "total_rows": total_rows,
                "total_cols": len(raw.columns),
                "raw_preview": raw_preview,
            })

    return result


def _sheet_meta(name: str, df: pd.DataFrame, suggested_header_row: int = 0) -> dict:
    col_info = []
    for col in df.columns:
        col_info.append({
            "name": col,
            "dtype": str(df[col].dtype),
            "nulls": int(df[col].isna().sum()),
            "sample": [str(v) for v in df[col].dropna().head(3).tolist()],
        })
    meta = {
        "sheet": name,
        "rows": len(df),
        "columns": len(df.columns),
        "column_info": col_info,
        "preview": df.head(15).to_markdown(index=False),
    }
    if suggested_header_row > 0:
        meta["suggested_header_row"] = suggested_header_row
    return meta


def read_excel(file_path: str) -> str:
    xl = pd.ExcelFile(file_path)
    parts = []
    for sheet_name in xl.sheet_names:
        df = xl.parse(sheet_name)
        # flatten multi-level columns nếu có
        if isinstance(df.columns, pd.MultiIndex):
            df.columns = [" > ".join(str(c) for c in col).strip() for col in df.columns]
        df.columns = [str(c) for c in df.columns]
        parts.append(f"### Sheet: {sheet_name}\n{df.to_markdown(index=False)}")
    return "\n\n".join(parts)


def read_docx(file_path: str) -> str:
    doc = Document(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if rows:
            df = pd.DataFrame(rows[1:], columns=rows[0])
            parts.append(df.to_markdown(index=False))
    return "\n\n".join(parts)


def read_markdown(file_path: str) -> str:
    return Path(file_path).read_text(encoding="utf-8")


def read_file(file_path: str) -> tuple[str, str]:
    """Tự detect loại file và đọc. Trả về (content, file_type)."""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return read_pdf(file_path), "pdf"
    elif ext in (".xlsx", ".xls"):
        return read_excel(file_path), "xlsx"
    elif ext == ".csv":
        return read_csv(file_path), "xlsx"
    elif ext == ".docx":
        return read_docx(file_path), "docx"
    elif ext == ".md":
        return read_markdown(file_path), "md"
    else:
        raise ValueError(f"Không hỗ trợ định dạng: {ext}")
